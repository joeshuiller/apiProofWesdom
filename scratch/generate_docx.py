import json
import os
import sys

# Importar docx (se instalará en el entorno virtual antes de ejecutar)
try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Error: docx library not found. Please install python-docx.")
    sys.exit(1)

# Paleta de colores ejecutivos
HEX_PRIMARY = "0F4C81"     # Azul clásico (Primary)
HEX_SECONDARY = "5B92E5"   # Azul claro (Secondary)
HEX_TEXT = "333333"        # Gris oscuro para texto
HEX_LIGHT_BG = "F4F6F9"    # Gris claro para fondo de cabeceras de tabla
HEX_BORDER = "CCCCCC"      # Borde gris claro

COLOR_PRIMARY = RGBColor(15, 76, 129)
COLOR_SECONDARY = RGBColor(91, 146, 229)
COLOR_TEXT = RGBColor(51, 51, 51)
COLOR_MUTED = RGBColor(120, 120, 120)

def set_cell_background(cell, fill_hex):
    """Establece el color de fondo de una celda de tabla."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Establece los márgenes internos (padding) de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Aplica bordes finos y limpios a toda la tabla."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_styled_heading(doc, text, level):
    """Agrega un encabezado estilizado con la paleta de colores del proyecto."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        
        # Línea de borde inferior para Heading 1
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12') # ~1.5pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), HEX_PRIMARY)
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
        
    elif level == 2:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        
    elif level == 3:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        
    return p

def extract_comment(comment_obj):
    """Extrae el texto del comentario en formato legible."""
    if not comment_obj:
        return ""
    text_parts = []
    
    # Extraer el summary principal
    summary = comment_obj.get("summary", [])
    for part in summary:
        text_parts.append(part.get("text", ""))
        
    # Extraer block tags adicionales si existen (e.g. @deprecated, @example)
    block_tags = comment_obj.get("blockTags", [])
    for tag in block_tags:
        tag_name = tag.get("tag", "").replace("@", "").capitalize()
        tag_content = []
        for part in tag.get("content", []):
            tag_content.append(part.get("text", ""))
        if tag_content:
            text_parts.append(f"\n*{tag_name}:* {' '.join(tag_content)}")
            
    return "".join(text_parts).strip()

def get_type_string(type_obj):
    """Convierte un objeto de tipo de TypeDoc en una cadena legible de TypeScript."""
    if not type_obj:
        return "any"
    
    t_type = type_obj.get("type")
    
    if t_type == "intrinsic":
        return type_obj.get("name", "any")
    elif t_type == "reference":
        name = type_obj.get("name", "any")
        type_args = type_obj.get("typeArguments", [])
        if type_args:
            args_str = ", ".join([get_type_string(arg) for arg in type_args])
            return f"{name}<{args_str}>"
        return name
    elif t_type == "union":
        types = type_obj.get("types", [])
        return " | ".join([get_type_string(t) for t in types])
    elif t_type == "array":
        elem_type = type_obj.get("elementType")
        return f"{get_type_string(elem_type)}[]"
    elif t_type == "reflection" and "declaration" in type_obj:
        # Objeto literal o firma de función
        dec = type_obj["declaration"]
        if "signatures" in dec:
            # Es una firma de función (callback)
            sig = dec["signatures"][0]
            params = []
            for p in sig.get("parameters", []):
                p_name = p.get("name", "")
                p_type = get_type_string(p.get("type"))
                params.append(f"{p_name}: {p_type}")
            ret = get_type_string(sig.get("type"))
            return f"({', '.join(params)}) => {ret}"
        return "object"
    return "any"

def parse_reflection(reflection):
    """Mapea una reflexión de TypeDoc (clase, interfaz, función, etc.) a un formato simplificado."""
    name = reflection.get("name", "")
    kind_id = reflection.get("kind", 0)
    
    kind_str = "Unknown"
    if kind_id == 128:
        kind_str = "Clase"
    elif kind_id == 256:
        kind_str = "Interfaz"
    elif kind_id == 64:
        kind_str = "Función"
    elif kind_id == 32:
        kind_str = "Variable"
    
    description = extract_comment(reflection.get("comment"))
    
    properties = []
    methods = []
    
    children = reflection.get("children", [])
    for child in children:
        child_kind = child.get("kind", 0)
        child_name = child.get("name", "")
        
        # Propiedad / Variable miembro (Kind = 1024)
        if child_kind == 1024:
            p_type = get_type_string(child.get("type"))
            p_desc = extract_comment(child.get("comment"))
            flags = child.get("flags", {})
            p_scope = "public"
            if flags.get("isPrivate"): p_scope = "private"
            elif flags.get("isProtected"): p_scope = "protected"
            
            properties.append({
                "name": child_name,
                "type": p_type,
                "description": p_desc,
                "scope": p_scope
            })
            
        # Método (Kind = 2048) o Constructor (Kind = 512)
        elif child_kind in [2048, 512]:
            signatures = child.get("signatures", [])
            for sig in signatures:
                sig_name = sig.get("name", child_name)
                sig_desc = extract_comment(sig.get("comment"))
                
                params = []
                for p in sig.get("parameters", []):
                    p_name = p.get("name", "")
                    p_type = get_type_string(p.get("type"))
                    p_desc = extract_comment(p.get("comment"))
                    params.append({
                        "name": p_name,
                        "type": p_type,
                        "description": p_desc
                    })
                
                ret_type = get_type_string(sig.get("type"))
                
                flags = child.get("flags", {})
                m_scope = "public"
                if flags.get("isPrivate"): m_scope = "private"
                elif flags.get("isProtected"): m_scope = "protected"
                
                methods.append({
                    "name": sig_name,
                    "description": sig_desc,
                    "parameters": params,
                    "return_type": ret_type,
                    "scope": m_scope,
                    "is_constructor": child_kind == 512
                })
                
    return {
        "name": name,
        "kind": kind_str,
        "description": description,
        "properties": properties,
        "methods": methods
    }

def collect_all_symbols(node, source_file=None):
    """Recolecta de forma recursiva todos los símbolos documentables del AST JSON."""
    symbols = []
    
    sources = node.get("sources", [])
    if sources:
        source_file = sources[0].get("fileName", source_file)
        
    kind = node.get("kind", 0)
    if kind in [128, 256]:
        parsed = parse_reflection(node)
        parsed["file"] = source_file
        symbols.append(parsed)
        
    for child in node.get("children", []):
        symbols.extend(collect_all_symbols(child, source_file))
        
    return symbols

def main():
    json_path = "scratch_doc.json"
    if not os.path.exists(json_path):
        print(f"Error: JSON file {json_path} does not exist.")
        sys.exit(1)
        
    print(f"Leyendo archivo AST JSON: {json_path}...")
    with open(json_path, "r", encoding="utf-8") as f:
        doc_data = json.load(f)
        
    all_symbols = collect_all_symbols(doc_data)
    print(f"Se encontraron {len(all_symbols)} clases/interfaces documentables.")
    
    categorized = {
        "Domain Layer (Dominio)": [],
        "Application Layer (Aplicación)": [],
        "Infrastructure Layer (Infraestructura)": [],
        "Core / Utils Layer (Core)": [],
        "Otras Declaraciones": []
    }
    
    for sym in all_symbols:
        filepath = sym.get("file", "")
        if "src/domain" in filepath:
            categorized["Domain Layer (Dominio)"].append(sym)
        elif "src/application" in filepath:
            categorized["Application Layer (Aplicación)"].append(sym)
        elif "src/infrastructure" in filepath:
            categorized["Infrastructure Layer (Infraestructura)"].append(sym)
        elif "src/core" in filepath:
            categorized["Core / Utils Layer (Core)"].append(sym)
        else:
            categorized["Otras Declaraciones"].append(sym)
            
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    for _ in range(3):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("DOCUMENTACIÓN TÉCNICA DE CÓDIGO\n")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Arquitectura Limpia (DDD) y Especificaciones de API en apiProofWesdom")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_SECONDARY
    
    for _ in range(8):
        doc.add_paragraph()
        
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        "Autor: Antigravity AI & Team Softsaenz S.A.S\n"
        "Nivel: Especificación Senior\n"
        "Lenguaje: TypeScript / Node.js\n"
        "Fecha de Generación: Julio 2026\n"
        "Versión: 1.0.0"
    )
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = COLOR_TEXT
    
    doc.add_page_break()
    
    add_styled_heading(doc, "1. Introducción y Arquitectura del Sistema", level=1)
    
    intro_text = (
        "El proyecto apiProofWesdom es un sistema de backend diseñado bajo las pautas de "
        "Domain-Driven Design (DDD) y Arquitectura Limpia (Clean Architecture). Esto garantiza una separación "
        "estricta de responsabilidades, facilitando el mantenimiento, la escalabilidad y la inyección de dependencias "
        "de manera desacoplada utilizando InversifyJS como contenedor IoC.\n\n"
        "La base de persistencia utiliza PostgreSQL 15 integrado a través de la librería TypeORM, manejando "
        "relaciones complejas y persistencia de cuentas, registros de auditoría de seguridad y transacciones de billeteras."
    )
    doc.add_paragraph(intro_text).style.font.name = 'Arial'
    
    sec_num = 2
    for layer_name, symbols in categorized.items():
        if not symbols:
            continue
            
        add_styled_heading(doc, f"{sec_num}. {layer_name}", level=1)
        sec_num += 1
        
        layer_desc = ""
        if "Domain" in layer_name:
            layer_desc = ("Contiene la lógica de negocio pura y el núcleo del dominio. "
                          "No depende de ningún framework externo ni detalles de bases de datos. "
                          "Define las entidades fundamentales y los contratos de almacenamiento.")
        elif "Application" in layer_name:
            layer_desc = ("Contiene los casos de uso que coordinan y orquestan el flujo de datos. "
                          "Implementa las reglas de aplicación específicas y define los DTOs de transferencia.")
        elif "Infrastructure" in layer_name:
            layer_desc = ("Implementa los detalles tecnológicos concretos como controladores Express, "
                          "rutas HTTP, adaptadores de Socket.IO, persistencia de base de datos (TypeORM) y middlewares.")
        elif "Core" in layer_name:
            layer_desc = ("Ofrece utilidades transversales (Cross-Cutting Concerns) como utilidades de "
                          "criptografía (RSA, AES), hashing de contraseñas con Argon2 y utilidades de estructuración.")
            
        if layer_desc:
            doc.add_paragraph(layer_desc)
            
        for sym in symbols:
            name = sym["name"]
            kind = sym["kind"]
            desc = sym["description"] or "Sin descripción detallada disponible."
            file = sym["file"] or ""
            
            add_styled_heading(doc, f"{kind}: {name}", level=2)
            
            meta_file_p = doc.add_paragraph()
            meta_file_run = meta_file_p.add_run(f"Ubicación del Archivo: {file}")
            meta_file_run.font.italic = True
            meta_file_run.font.size = Pt(9.5)
            meta_file_run.font.color.rgb = COLOR_MUTED
            
            doc.add_paragraph(desc)
            
            if sym["properties"]:
                add_styled_heading(doc, f"Propiedades y Atributos de {name}", level=3)
                
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Propiedad / Variable'
                hdr_cells[1].text = 'Tipo de Dato (TS)'
                hdr_cells[2].text = 'Descripción / Ámbito'
                
                for cell in hdr_cells:
                    set_cell_background(cell, HEX_PRIMARY)
                    set_cell_margins(cell, top=120, bottom=120)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            
                for prop in sym["properties"]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f"{prop['name']} ({prop['scope']})"
                    row_cells[1].text = prop['type']
                    row_cells[2].text = prop['description'] or "No especificado."
                    
                    for cell in row_cells:
                        set_cell_margins(cell, top=80, bottom=80)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9.5)
                                run.font.color.rgb = COLOR_TEXT
                                
                doc.add_paragraph()
                
            if sym["methods"]:
                add_styled_heading(doc, f"Métodos y Funciones de {name}", level=3)
                
                for m in sym["methods"]:
                    m_name = m["name"]
                    m_scope = m["scope"]
                    m_desc = m["description"] or "Sin descripción detallada."
                    m_ret = m["return_type"]
                    
                    param_strs = []
                    for p in m["parameters"]:
                        param_strs.append(f"{p['name']}: {p['type']}")
                    signature_str = f"{m_scope} {m_name}({', '.join(param_strs)}): {m_ret}"
                    
                    sig_p = doc.add_paragraph()
                    sig_p.paragraph_format.left_indent = Inches(0.2)
                    sig_run = sig_p.add_run(signature_str)
                    sig_run.font.bold = True
                    sig_run.font.name = 'Courier New'
                    sig_run.font.size = Pt(9.5)
                    sig_run.font.color.rgb = COLOR_SECONDARY
                    
                    desc_p = doc.add_paragraph()
                    desc_p.paragraph_format.left_indent = Inches(0.2)
                    desc_run = desc_p.add_run(m_desc)
                    desc_run.font.name = 'Arial'
                    desc_run.font.size = Pt(9.5)
                    desc_run.font.color.rgb = COLOR_TEXT
                    
                    if m["parameters"]:
                        p_table = doc.add_table(rows=1, cols=3)
                        p_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        set_table_borders(p_table)
                        
                        p_hdr = p_table.rows[0].cells
                        p_hdr[0].text = 'Parámetro'
                        p_hdr[1].text = 'Tipo'
                        p_hdr[2].text = 'Descripción'
                        
                        for cell in p_hdr:
                            set_cell_background(cell, "E0E6ED")
                            set_cell_margins(cell, top=80, bottom=80)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = COLOR_PRIMARY
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(9)
                                    
                        for p in m["parameters"]:
                            p_row = p_table.add_row().cells
                            p_row[0].text = p['name']
                            p_row[1].text = p['type']
                            p_row[2].text = p['description'] or "No especificado."
                            
                            for cell in p_row:
                                set_cell_margins(cell, top=60, bottom=60)
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(9)
                                        run.font.color.rgb = COLOR_TEXT
                                        
                        doc.add_paragraph()
                        
                doc.add_paragraph()
                
    output_filename = "DOCUMENTATION.docx"
    print(f"Guardando documento de Word como: {output_filename}...")
    doc.save(output_filename)
    print("¡Generación completada exitosamente!")

if __name__ == "__main__":
    main()
